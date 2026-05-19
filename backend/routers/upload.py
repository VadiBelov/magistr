from fastapi import APIRouter, UploadFile, File
from fastapi.responses import JSONResponse
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker
import os
import uuid
import config
from worker.tasks import transcribe_audio
from fastapi.responses import JSONResponse, HTMLResponse, FileResponse
import json

router = APIRouter()

engine = create_engine(config.DATABASE_URL)
SessionLocal = sessionmaker(bind=engine)


@router.post("/upload")
async def upload_audio(file: UploadFile = File(...)):
    if not file.filename.endswith(('.wav', '.mp3', '.ogg', '.flac', '.m4a', '.webm')):
        return JSONResponse({"error": "Только .wav, .mp3, .ogg, .flac, .m4a"}, status_code=400)

    file_id = str(uuid.uuid4())
    ext = file.filename.split('.')[-1]
    save_path = os.path.join(config.AUDIO_TEMP_DIR, f"{file_id}.{ext}")

    with open(save_path, "wb") as f:
        content = await file.read()
        f.write(content)

    file_size = len(content) / (1024 * 1024)

    # Отправляем задачу в Celery (фоновую)
    task = transcribe_audio.delay(save_path, file_id, file.filename)

    return {
        "status": "processing",
        "file_id": file_id,
        "filename": file.filename,
        "size_mb": round(file_size, 2),
        "task_id": task.id,
        "message": "Файл отправлен на распознавание. Проверьте статус по /result/{task_id}"
    }


@router.get("/result/{task_id}")
async def get_result(task_id: str):
    task = transcribe_audio.AsyncResult(task_id)

    if task.state == "PENDING":
        return {"status": "pending", "message": "Задача в очереди"}
    elif task.state == "STARTED":
        return {"status": "started", "message": "Идёт распознавание..."}
    elif task.state == "SUCCESS":
        return {"status": "done", "result": task.result}
    elif task.state == "FAILURE":
        return {"status": "error", "message": str(task.info)}
    else:
        return {"status": task.state}

@router.get("/", response_class=HTMLResponse)
async def index():
    from pathlib import Path
    html_path = Path(__file__).parent.parent.parent / "frontend" / "index.html"
    return html_path.read_text(encoding="utf-8")


@router.get("/export/{task_id}")
async def export_docx(task_id: str):
    from docx import Document

    task = transcribe_audio.AsyncResult(task_id)
    if task.state != "SUCCESS":
        return JSONResponse({"error": "Результат не готов"}, status_code=400)

    result = task.result
    doc = Document()
    doc.add_heading("Протокол встречи", level=1)
    doc.add_paragraph(f"Файл: {result.get('filename', '')}")
    doc.add_paragraph(f"Язык: {result.get('language', '')}")

    if result.get('summary'):
        doc.add_heading("Краткий итог", level=2)
        doc.add_paragraph(result['summary'])

    doc.add_heading("Расшифровка", level=2)
    for seg in result.get('segments', []):
        p = doc.add_paragraph()
        p.add_run(f"[{seg['start']}s-{seg['end']}s] ").bold = False
        p.add_run(f"{seg['speaker']}: ").bold = True
        p.add_run(seg['text'])

    export_dir = config.AUDIO_TEMP_DIR / "exports"
    export_dir.mkdir(exist_ok=True)
    export_path = export_dir / f"protocol_{task_id[:8]}.docx"
    doc.save(str(export_path))

    return FileResponse(
        str(export_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"protocol_{task_id[:8]}.docx"
    )


from fastapi import WebSocket, WebSocketDisconnect
from backend.ws_manager import manager
from faster_whisper import WhisperModel
import tempfile
import base64

live_model = None


def get_live_model():
    global live_model
    if live_model is None:
        print("⏳ Загрузка Whisper для live...")
        live_model = WhisperModel("small", device="cpu", compute_type="int8")
        print("✅ Live-модель готова")
    return live_model


from fastapi import WebSocket, WebSocketDisconnect
from backend.ws_manager import manager
from faster_whisper import WhisperModel
import tempfile
import base64

live_model = None

def get_live_model():
    global live_model
    if live_model is None:
        print("⏳ Загрузка Whisper для live...")
        live_model = WhisperModel("small", device="cpu", compute_type="int8")
        print("✅ Live-модель готова")
    return live_model

@router.websocket("/ws/room/{room_id}")
async def websocket_room(websocket: WebSocket, room_id: str):
    user_id = websocket.query_params.get("user_id", "anonymous")
    await websocket.accept()
    await manager.join_room(room_id, user_id, websocket)
    model = get_live_model()

    try:
        while True:
            data = await websocket.receive_json()

            # Обработка чата
            if "type" in data and data["type"] == "chat":
                await manager.broadcast_chat(room_id, user_id, data["text"])
                continue

            # Обработка аудио-чанков (если будут)
            audio_b64 = data.get("audio", "")
            if audio_b64:
                import os as os_module
                wav_bytes = base64.b64decode(audio_b64)

                if not hasattr(websocket, '_audio_buf'):
                    websocket._audio_buf = b""
                websocket._audio_buf += wav_bytes

                if len(websocket._audio_buf) > 96000:
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
                    tmp.write(websocket._audio_buf)
                    tmp.close()
                    websocket._audio_buf = b""

                    try:
                        segments, info = model.transcribe(tmp.name, beam_size=1, language="ru")
                        text = " ".join([s.text.strip() for s in segments])
                        if text.strip():
                            seg = {"speaker": user_id, "text": text.strip(), "start": 0, "end": 0}
                            manager.get_transcript(room_id).append(seg)
                            await manager.broadcast_transcript(room_id, seg)
                    except Exception as e:
                        print(f"Live error: {e}")

                    try: os_module.unlink(tmp.name)
                    except: pass

    except WebSocketDisconnect:
        pass
    except Exception as e:
        print(f"WS error: {e}")
    finally:
        await manager.leave_room(room_id, user_id, websocket)

@router.get("/room/{room_id}/transcript")
async def get_room_transcript(room_id: str):
    return {"room_id": room_id, "segments": manager.get_transcript(room_id)}

@router.post("/room/create")
async def create_room():
    room_id = str(uuid.uuid4())[:8]
    manager.create_room(room_id)
    return {"room_id": room_id}